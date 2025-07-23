import base64
import datetime
import io
import json
import logging
import asyncio
import time

import numpy as np
import pandas as pd
from aiohttp import ClientSession,ClientError

import mplfinance
import os
import pandas
import random
import re
import requests
import site
import subprocess
import tempfile
import threading
import tiktoken
import uuid
import w3lib.encoding
from bs4 import BeautifulSoup, Tag
from docx import Document
from flask import Blueprint, Response, request, send_file, make_response, jsonify
from flask_restful import reqparse
from newspaper import Article
from openpyxl.worksheet.worksheet import Worksheet
from pandas import DataFrame
from readabilipy.simple_json import plain_content, extract_text_blocks_as_plain_text
from readabilipy.utils import chdir

import env
import lmjj_agent
import paddle_ocr
import pdf_common_parse
import pdf_table_parse
import word_reader
import excel_to_json
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO
import get_file_path
import async_request
# 导入调度器模块
from scheduler import scheduler, TaskConfig

bp = Blueprint('controller', __name__)

logger = logging.getLogger(__name__)

tokeniser = tiktoken.get_encoding('cl100k_base')


@bp.route('/bing_news_search', methods=['POST', 'GET', ])
def bing_news_search():
    # https://learn.microsoft.com/en-us/bing/search-apis/bing-news-search/reference/endpoints

    parser = reqparse.RequestParser()
    parser.add_argument('article_count',
                        location='args', type=int, required=False, default=5)
    parser.add_argument('count',
                        location='args', type=int, required=False, default=0)
    parser.add_argument('freshness',
                        location='args', type=str, required=False, default='Week',
                        choices=['Day', 'Week', 'Month', ])
    parser.add_argument('q',
                        location='args', type=str, required=True)
    parser.add_argument('sort_by',
                        location='args', type=str, required=False, default='Date',
                        choices=['Date', 'Relevance', ])
    args = parser.parse_args()

    article_count = args.get('article_count')
    count = args.get('count')
    freshness = args.get('freshness')
    q = args.get('q')
    sort_by = args.get('sort_by')

    logger.info(
        f'必应新闻搜索'
        f'article_count[{article_count}]'
        f'q[{q}]'
    )

    cc = 'CN'
    mkt = 'zh-CN'
    set_lan = 'zh-hans'

    count = max(count, article_count * 10)
    count = min(count, 99)

    params = {
        'cc': cc,
        'count': count,
        'freshness': freshness,
        'mkt': mkt,
        "q": q,
        'setLang': set_lan,
        "sortBy": sort_by,
        "textDecorations": True,
        "textFormat": "HTML",
    }

    url = f'https://api.bing.microsoft.com/v7.0/news/search'
    response = requests.get(url, headers=env.bind_search_headers, params=params, verify=False)
    response.raise_for_status()
    response_json: dict = response.json()

    value: list[dict] = response_json.get('value')
    if not value:
        logger.error(
            f'必应新闻搜索-结果异常'
            f'[value为空]'
            f'[{response_json}]'
        )
        raise Exception('必应新闻搜索-结果异常')

    articles = []
    urls = []
    for item in value:
        logger.info(
            f'必应新闻搜索-news-item'
            f'[{item}]'
        )

        category = item.get('category')
        date_published = item.get('datePublished')
        description = item.get('description')
        image = item.get('image')
        name = item.get('name')
        provider = item.get('provider')
        url = item.get('url')

        article_json = extract_article(url)
        if not article_json:
            # 跳过无效的网页链接
            continue

        article = render_template(
            title=article_json.get('title'),
            authors=article_json.get('byline'),
            publish_date=article_json.get('date'),
            text=article_json.get('plain_text'),
        )
        articles.append(article)
        # 记录有效链接
        urls.append(url)

        if len(articles) >= article_count:
            # collected enough articles already
            break

    result = '\n\n\n'.join(articles)

    tokens = calc_tokens(result)
    logger.info(
        f'采集文章数量[{article_count}]'
        f'采集文章长度[{len(result)}]'
        f'预估消耗tokens[{tokens}]'
        f'原始文章链接[{urls}]'
    )

    return {
        'text': result,
        'urls': urls,
    }


def calc_tokens(text):
    tokens = len(tokeniser.encode(text))
    k_tokens = round(tokens / 1000.0)
    logger.info(f'预估tokens'
                f'[{tokens}]'
                f'[{k_tokens}k]')
    return tokens


@bp.route('/read_article', methods=['POST', 'GET', ])
def read_article() -> str:
    parser = reqparse.RequestParser()
    parser.add_argument('article_url',
                        location='args', type=str, required=True)
    args = parser.parse_args()

    article_url = args.get('article_url')

    # todo 六脉九剑貌似无法正确处理参数是 URL 的情况，临时变通一下
    article_json = extract_article(f'http://{article_url}')
    if not article_json:
        return ''

    article = render_template(
        title=article_json.get('title'),
        authors=article_json.get('byline'),
        publish_date=article_json.get('date'),
        text=article_json.get('plain_text'),
    )

    return article


def render_template(
        title,
        authors,
        publish_date,
        text: str,
):
    text = text.strip() if text else ''
    t = (
        f'TITLE: {title}\n'
        f'AUTHORS: {authors}\n'
        f'PUBLISH DATE: {publish_date}\n'
        f'TEXT:\n'
        f'\n'
        f'{text}'
    )
    return t


def extract_article(url: str) -> None | dict:
    h = {
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
        'Cache-Control': 'no-cache',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36',
    }
    try:
        resp = requests.get(url=url, headers=h, proxies=None, verify=False)
    except Exception as e:
        logger.error(
            f'访问网页-失败'
            f'[requests.get=>{url}]',
            exc_info=True
        )
        return None

    if resp.status_code != 200:
        logger.error(
            f'访问网页-失败'
            f'[status_code=>{resp.status_code}]'
            f'[{url}]'
        )
        return None

    logger.info(
        f'访问网页-成功'
        f'[{url}]'
    )

    declared = w3lib.encoding.html_body_declared_encoding(resp.text)
    if declared:
        resp.encoding = declared
    else:
        # default encoding
        resp.encoding = 'utf-8'

    html = resp.text

    content_digests = False
    node_indexes = False

    with tempfile.NamedTemporaryFile(encoding='utf-8', delete=False, mode='w') as temp_html:
        temp_html.write(html)
        temp_html.close()

    html_path = temp_html.name
    article_json_path = f'{html_path}.article.json'
    logger.info(
        f'html_path[{html_path}]'
        f'article_json_path[{article_json_path}]'
    )

    def find_module_path(module_name):
        for package_path in site.getsitepackages():
            potential_path = os.path.join(package_path, module_name)
            if os.path.exists(potential_path):
                return potential_path
        raise Exception()

    js_dir = os.path.join(find_module_path('readabilipy'), 'javascript')
    with chdir(js_dir):
        try:
            subprocess.check_call(["node", "ExtractArticle.js", "-i", html_path, "-o", article_json_path])
        except Exception as e:
            logger.error(
                f'提取网页-失败'
                f'[ExtractArticle.js]'
                f'[{url}]'
                f'[{html}]',
                exc_info=True
            )
            return None

    with open(article_json_path, encoding='utf-8') as f:
        input_json = json.loads(f.read())

    os.remove(html_path)
    os.remove(article_json_path)

    article_json = {
        "title": None,
        "byline": None,
        "date": None,
        "content": None,
        "plain_content": None,
        "plain_text": None,

        "text_content": None,
        "length": None,
        "excerpt": None,
    }
    if input_json:
        if "title" in input_json and input_json["title"]:
            article_json["title"] = input_json["title"]
        if "byline" in input_json and input_json["byline"]:
            article_json["byline"] = input_json["byline"]
        if "date" in input_json and input_json["date"]:
            article_json["date"] = input_json["date"]
        if "content" in input_json and input_json["content"]:
            article_json["content"] = input_json["content"]
            article_json["plain_content"] = plain_content(article_json["content"], content_digests, node_indexes)
            article_json["plain_text"] = extract_text_blocks_as_plain_text(article_json["plain_content"])

        # 参考 DI-FY
        if "textContent" in input_json and input_json["textContent"]:
            article_json["plain_text"] = input_json["textContent"]
            article_json["plain_text"] = re.sub(r'\n\s*\n', '\n', article_json["plain_text"])

        if 'textContent' in input_json and input_json['textContent']:
            article_json['text_content'] = input_json['textContent']
        if 'length' in input_json and input_json['length'] is not None:
            article_json['length'] = input_json['length']
        if 'excerpt' in input_json and input_json['excerpt'] is not None:
            article_json['excerpt'] = input_json['excerpt']
    else:
        logger.error(
            f'提取网页-失败'
            f'[ExtractArticle.js]'
            f'[input_json 为空]'
            f'[{url}]'
            f'[{html}]'
        )

    if not article_json['plain_text'] or not article_json['plain_text'].strip():
        logger.error(
            f'提取网页-失败'
            f'[ExtractArticle.js]'
            f'[plain_text 为空]'
            f'[{input_json}]'
            f'[{url}]'
            f'[{html}]'
        )
        # todo fail-over to newspaper3k
        # a = extract_article_by_newspaper3k(url)
        return None

    return article_json


def extract_article_by_newspaper3k(url) -> Article:
    a = Article(url)
    a.download()
    a.parse()
    return a


article_staging_dict: dict = {}


@bp.route('/article_staging', methods=['POST', ])
def article_staging():
    parser = reqparse.RequestParser()
    parser.add_argument('article_bag_id',
                        location='json', type=str, required=True)
    parser.add_argument('article_content',
                        location='json', type=str, required=False)
    parser.add_argument('action',
                        location='json', type=str, required=False, default='put',
                        choices=['put', 'get', 'del', ])
    args = parser.parse_args()

    article_bag_id = args.get('article_bag_id')
    article_content = args.get('article_content')
    action = args.get('action')

    # todo 不太可能出现并发冲突，所以不用加锁

    if action == 'put':
        if not article_content:
            raise Exception('[article_content]不能为空')

        article_staging_dict.setdefault(article_bag_id, []).append(article_content)
        return {}
    elif action == 'get':
        return {
            'data': article_staging_dict.get(article_bag_id, [])
        }
    elif action == 'del':
        article_staging_dict.pop(article_bag_id, '')
        return {}
    else:
        raise Exception()


def load_headers(path: str):
    headers = {}
    with open(path, encoding='utf-8', mode='r') as f:
        for line in f.readlines():
            idx = line.find(':')
            key = line[:idx]
            val = line[(idx + 2):]
            headers[key] = val.strip()
    return headers


def _sge_graph_daily_quotation_json() -> dict:
    """
    行情走势-JSON
    :return:
    """
    # https://www.sge.com.cn/sjzx/mrhq

    parser = reqparse.RequestParser()
    parser.add_argument('inst_id',
                        location='args', type=str, required=False,
                        default='Au(T+D)')
    parser.add_argument('limit',
                        location='args', type=int, required=False,
                        default=30)
    args = parser.parse_args()

    inst_id = args.get('inst_id')
    limit = args.get('limit')

    data = {
        'instid': inst_id,
    }
    h = load_headers('headers/sge/headers-post.txt')
    url = f'https://www.sge.com.cn/graph/Dailyhq'
    resp = requests.post(url=url, headers=h, data=data, proxies=None, verify=False)

    resp_json: dict = resp.json()

    time_list = resp_json.get('time')

    # 只保留最近 limit 个交易日的数据
    time_list = time_list[-limit:]

    # some_date
    # open_price
    # close_price
    # lowest_price
    # highest_price
    return {
        'inst_id': inst_id,
        'limit': limit,
        'time_list': time_list,
    }


def sge_graph_daily_quotation_plot():
    j = _sge_graph_daily_quotation_json()
    inst_id = j['inst_id']
    limit = j['limit']
    time_list = j['time_list']

    df: DataFrame = pandas.DataFrame(data=time_list, columns=['Date', 'Open', 'Close', 'Low', 'High'])
    df['Date'] = pandas.to_datetime(df['Date'])
    df.set_index('Date', inplace=True)

    logger.info(
        f'行情走势-PLOT'
        f'[{df}]'
    )
    buf = io.BytesIO()

    # https://github.com/matplotlib/mplfinance/wiki/Mplfinance-Style-Sheets-Reference
    dpi = 300
    width = 1920
    height = 856
    # todo 解决标题/label中文乱码
    mplfinance.plot(
        data=df, type='candle', style='charles',
        title='Au(T+D)', xlabel='Date', ylabel='Price',
        savefig=dict(fname=buf, dpi=300, ),
        figsize=(width / dpi, height / dpi),
        scale_padding={'left': 1, 'top': 5, 'right': 1, 'bottom': 1},
        tight_layout=True,
    )
    buf.seek(0)
    return buf


@bp.route('/sge_graph_daily_quotation_img', methods=['POST', 'GET', ])
def sge_graph_daily_quotation_img():
    buf = sge_graph_daily_quotation_plot()
    return Response(buf, mimetype='image/png')


@bp.route('/sge_graph_daily_quotation_img_base64', methods=['POST', 'GET', ])
def sge_graph_daily_quotation_img_base64():
    buf = sge_graph_daily_quotation_plot()
    b64 = base64.b64encode(buf.read()).decode('utf-8')
    return b64


@bp.route('/sge_graph_daily_quotation_json', methods=['POST', 'GET', ])
def sge_graph_daily_quotation_json():
    """
    行情走势-JSON
    :return:
    """
    j = _sge_graph_daily_quotation_json()
    time_list = j['time_list']
    return time_list


@bp.route('/sge_graph_daily_quotation', methods=['POST', 'GET', ])
def sge_graph_daily_quotation():
    """
    行情走势-TEXT
    :return:
    """
    j = _sge_graph_daily_quotation_json()
    time_list = j['time_list']

    rows = []
    for when, open_p, close_p, lowest_p, highest_p in time_list:
        when: str = when
        row = f'|{when}|{open_p}|{close_p}|{lowest_p}|{highest_p}|'
        rows.append(row)

    row_str = '\n'.join(rows)
    result = (
        f'| 日期 | 开盘价 | 收盘价 | 最低价 | 最高价 |\n'
        f'| ---- | ---- | ---- | ---- | ---- |\n'
        f'{row_str}'
    )

    calc_tokens(result)

    return result


def _sge_daily_quotation_new() -> DataFrame:
    """
    历史行情数据-每日行情
    :return:
    """
    # https://www.sge.com.cn/sjzx/quotation_daily_new

    today = datetime.datetime.now().date()
    early = today - datetime.timedelta(days=30)

    today = str(today)
    early = str(early)

    parser = reqparse.RequestParser()
    parser.add_argument('start_date',
                        location='args', type=str, required=False,
                        default=early)
    parser.add_argument('end_date',
                        location='args', type=str, required=False,
                        default=today)
    parser.add_argument('inst_ids',
                        location='args', type=str, required=False,
                        default='Au(T+D)')
    args = parser.parse_args()
    start_date = args.get('start_date')
    end_date = args.get('end_date')
    inst_ids = args.get('inst_ids')

    params = {
        'start_date': start_date,
        'end_date': end_date,
        'inst_ids': inst_ids,
        # todo 数据量多了会涉及分页
        'p': 1,
    }
    h = load_headers('./headers/sge/headers-get.txt')
    url = f'https://www.sge.com.cn/sjzx/quotation_daily_new'
    resp = requests.get(url=url, headers=h, params=params, proxies=None, verify=False)
    resp_text = resp.text
    # 日期
    # 合约
    # 开盘价
    # 最高价
    # 最低价
    # 收盘价
    # 涨跌（元）
    # 涨跌幅
    # 加权平均价
    # 成交量（kg）
    # 成交金额（元）
    # 市场持仓（手）
    # 交收方向
    # 交收量（手）
    df: DataFrame = pandas.read_html(str(resp_text))[0]
    logger.info(f'请求[历史行情数据-每日行情][{params}]=>{df}')

    return df


@bp.route('/sge_daily_quotation_new_json')
def sge_daily_quotation_new_json():
    """
    历史行情数据-每日行情-JSON
    :return:
    """
    df = _sge_daily_quotation_new()
    data_list = []
    if df.size:
        data_list = df.values.tolist()

    return data_list


@bp.route('/sge_daily_quotation_new')
def sge_daily_quotation_new():
    """
    历史行情数据-每日行情-TEXT
    :return:
    """
    df = _sge_daily_quotation_new()

    rows = []
    if df.size:
        data_list = df.values.tolist()
        for items in data_list:
            row = [str(item) for item in items]
            row = '|'.join(row)
            rows.append(row)

    cols = df.columns.tolist()
    head = '|'.join([str(item) for item in cols])
    line = '|'.join([' ---- ' for _ in cols])

    data = '\n'.join([f'|{row}|' for row in rows])
    result = (
        f'|{head}|\n'
        f'|{line}|\n'
        f'{data}'
    )

    calc_tokens(result)

    return result


@bp.route('/sge_daily_quotation_new_raw')
def sge_daily_quotation_new_raw():
    """
    历史行情数据-每日行情-TEXT-非markdown格式
    :return:
    """
    df = _sge_daily_quotation_new()

    return str(df)


@bp.route('/chicken_soup')
def chicken_soup():
    # https://www.xinlingjitang.net/
    url = 'https://www.xinlingjitang.net/'

    h = {
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
        'Accept-Encoding': 'gzip',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
        'Cache-Control': 'no-cache',
        'Pragma': 'no-cache',
        'Sec-Ch-Ua': '"Chromium";v="122", "Not(A:Brand";v="24", "YaBrowser";v="24.4", "Yowser";v="2.5"',
        'Sec-Ch-Ua-Mobile': '?0',
        'Sec-Ch-Ua-Platform': '"Windows"',
        'Sec-Fetch-Dest': 'document',
        'Sec-Fetch-Mode': 'navigate',
        'Sec-Fetch-Site': 'none',
        'Sec-Fetch-User': '?1',
        'Upgrade-Insecure-Requests': '1',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36',
    }
    resp = requests.get(url=url, headers=h, verify=False)

    # todo 需要抽取单独的方法
    declared = w3lib.encoding.html_body_declared_encoding(resp.text)
    if declared:
        resp.encoding = declared
    else:
        # default encoding
        resp.encoding = 'utf-8'

    html = resp.text

    soup = BeautifulSoup(html)
    contents = soup.find('div', attrs={'id': 'index-content', })
    p_tags: list[Tag] = [p for p in contents]

    if p_tags:
        for _ in range(0, 5):
            idx = random.randint(0, len(p_tags))
            p_text = p_tags[idx].text
            if p_text:
                logger.info(f'本次抓取到的心灵鸡汤[{p_text}]')
                return p_text

    logger.error(f'[心灵鸡汤抓取失败]=>{p_tags}\n'
                 f'[抓取到的网页源码]=>{html}')
    return '心灵鸡汤抓取失败'


@bp.route('/write_excel', methods=['POST', 'GET', ])
def write_excel():
    req_text = request.data.decode('utf-8')
    logger.info(f'write excel for => {req_text}')
    try:
        try:
            data: dict = json.loads(req_text)
        except Exception:
            logger.error(f'请求参数不是正确的JSON=>{req_text}', exc_info=True)
            return {
                'error': '请求参数不是正确的JSON',
            }

        if not data:
            return {
                'error': 'JSON不能为空',
            }
        app_code = request.args.get('app_code')

        f_name = './file/营销活动申请表空白模板.xlsx'
        f_path = os.path.abspath(f_name)
        from openpyxl import load_workbook

        wb_template = load_workbook(filename=f_path)
        ws: Worksheet = wb_template.get_sheet_by_name(name='消费类活动申请表')

        for position, val in data.items():
            try:
                ws[position] = str(val)
            except Exception:
                logger.error(f'单元格填写失败[{position}][{val}]', exc_info=True)
                return {
                    'error': f'单元格填写失败[{position}][{val}]',
                }

        tmp_filename = f'营销活动申请表-{str(uuid.uuid4())}.xlsx'
        tmp_filepath = f'{tempfile.gettempdir()}/{tmp_filename}'
        try:
            wb_template.save(filename=tmp_filepath)
        except Exception:
            logger.error(f'临时文件生成失败[{tmp_filepath}]', exc_info=True)
            return {
                'error': f'临时文件生成失败',
            }

        # fixme 先固定使用70服务上传文件
        remote_addr: str = '172.16.1.70'

        result = lmjj_agent.upload_file(app_code=app_code, file=tmp_filepath, remote_addr=remote_addr)
        result['error'] = ''
        return result

    except Exception:
        logger.error(f'Excel填写失败', exc_info=True)
        return {
            'error': 'Excel填写失败',
        }


@bp.route('/write_excel_v2', methods=['POST', 'GET', ])
def write_excel_v2():
    parser = reqparse.RequestParser()
    parser.add_argument('app_code',
                        location='args', type=str, required=True, )
    parser.add_argument('json_data',
                        location='json', type=dict, required=True, )
    parser.add_argument('from_addr',
                        location='json', type=str, required=False)
    args = parser.parse_args()

    app_code = args.get('app_code')
    json_data: dict = args.get('json_data')
    from_addr = args.get('from_addr')

    try:
        f_name = './file/营销活动申请表空白模板.xlsx'
        f_path = os.path.abspath(f_name)
        from openpyxl import load_workbook

        wb_template = load_workbook(filename=f_path)
        ws: Worksheet = wb_template.get_sheet_by_name(name='消费类活动申请表')

        for position, val in json_data.items():
            try:
                ws[position] = str(val)
            except Exception:
                logger.error(f'单元格填写失败[{position}][{val}]', exc_info=True)
                return {
                    'error': f'单元格填写失败[{position}][{val}]',
                }

        tmp_filename = f'营销活动申请表-{str(uuid.uuid4())}.xlsx'
        tmp_filepath = f'{tempfile.gettempdir()}/{tmp_filename}'
        try:
            wb_template.save(filename=tmp_filepath)
        except Exception:
            logger.error(f'临时文件生成失败[{tmp_filepath}]', exc_info=True)
            return {
                'error': f'临时文件生成失败',
            }

        result = lmjj_agent.upload_file(app_code=app_code, file=tmp_filepath, remote_addr=from_addr)
        result['error'] = ''
        return result

    except Exception:
        logger.error(f'Excel填写失败', exc_info=True)
        return {
            'error': 'Excel填写失败',
        }


@bp.route('/read_word', methods=['POST', 'GET', ])
def read_word():
    parser = reqparse.RequestParser()
    parser.add_argument('file_url',
                        location='json', type=str, required=True, )
    args = parser.parse_args()

    file_url = args.get('file_url')
    resp = requests.get(url=file_url, verify=False)
    try:
        doc: Document = Document(io.BytesIO(resp.content))
    except Exception as e:
        logger.error(f'文件解析失败，可能不是word文档[{file_url}][{str(e)}]')
        return '文件解析失败，可能不是word文档'

    logger.info(f'文件解析成功[{file_url}]')
    all_text = []
    for para in doc.paragraphs:
        text = para.text
        if text:
            all_text.append(text)
    # join all texts
    result = '\n'.join(all_text)
    return result


@bp.route('/read_word_v2', methods=['POST', 'GET', ])
def read_word_v2():
    parser = reqparse.RequestParser()
    parser.add_argument('app_code',
                        location='json', type=str, required=True, )
    parser.add_argument('file_info',
                        location='json', type=str, required=True, )
    parser.add_argument('from_addr',
                        location='json', type=str, required=False)
    args = parser.parse_args()

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')

    logger.info(f'read_word_v2[{app_code}][{file_info}][{from_addr}]')

    head = file_info.find("related_id='")
    file_info = file_info[head + len("related_id='"):]
    tail = file_info.find("'")
    file_id = file_info[:tail]

    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'read_word_v2[{app_code}][{file_info}][{remote_addr}]')

    return lmjj_agent.file_preview(app_code, file_id=file_id, remote_addr=remote_addr)


@bp.route('/parse_pdf_table', methods=['POST', 'GET', ])
def parse_pdf_table():
    parser = reqparse.RequestParser()
    parser.add_argument('app_code',
                        location='json', type=str, required=True, )
    parser.add_argument('file_info',
                        location='json', type=str, required=True, )
    parser.add_argument('from_addr',
                        location='json', type=str, required=False)
    args = parser.parse_args()

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')

    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'parse_pdf_table[{app_code}][{file_info}][{remote_addr}]')

    head = file_info.find("tenant_id='") + len("tenant_id='")
    tail = head + 36
    tenant_id = file_info[head:tail]

    head = file_info.find("related_id='") + len("related_id='")
    tail = head + 36
    file_id = file_info[head:tail]

    file_path = lmjj_agent.file_download_by_tenant(
        app_code=app_code,
        file_id=file_id,
        tenant_id=tenant_id,
        remote_addr=remote_addr
    )

    return {'text': pdf_table_parse.extract_tables(file_path)}


@bp.route('/read_word_v3', methods=['POST', 'GET', ])
def read_word_v3():
    req_text = request.data.decode('utf-8')
    logger.info(
        f'read_word_v3'
        f'[{req_text}]'
    )
    try:
        args: dict = json.loads(req_text)
    except Exception:
        logger.error(
            f'read_word_v3'
            f'[请求参数不是正确的JSON]',
            exc_info=True
        )
        return {
            'error': '请求参数不是正确的JSON',
            'text': '',
        }

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')

    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'read_word_v3[{app_code}][{file_info}][{remote_addr}]')

    head = file_info.find("tenant_id='") + len("tenant_id='")
    tail = head + 36
    tenant_id = file_info[head:tail]

    head = file_info.find("related_id='") + len("related_id='")
    tail = head + 36
    file_id = file_info[head:tail]

    file_path = lmjj_agent.file_download_by_tenant(
        app_code=app_code,
        file_id=file_id,
        tenant_id=tenant_id,
        remote_addr=remote_addr
    )
    doc: Document = Document(file_path)
    all_text = []
    for para in doc.paragraphs:
        text = para.text
        if text:
            all_text.append(text)
    # join all texts
    return {'text': '\n'.join(all_text)}


@bp.route('/read_word_v4', methods=['POST', 'GET', ])
def read_word_v4():
    req_text = request.data.decode('utf-8')

    logger.info(
        f'read_word_v4'
        f'[{req_text}]'
    )
    try:
        args: dict = json.loads(req_text)
    except Exception:
        logger.error(
            f'read_word_v4'
            f'[请求参数不是正确的JSON]',
            exc_info=True
        )
        return {
            'error': '请求参数不是正确的JSON',
            'text': '',
        }

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')

    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'read_word_v4[{app_code}][{file_info}][{remote_addr}]')

    # 目前最多只能上传[6]个文件
    # 预分配参数防止工作流crash
    data = {
        'text_1': '',
        'text_2': '',
        'text_3': '',
        'text_4': '',
        'text_5': '',
        'text_6': '',
    }
    idx = 1
    while True:
        head = file_info.find("tenant_id='")
        if head < 0:
            break
        head += + len("tenant_id='")
        tail = head + 36
        tenant_id = file_info[head:tail]

        head = file_info.find("related_id='")
        head += + len("related_id='")
        tail = head + 36
        file_id = file_info[head:tail]

        file_path = lmjj_agent.file_download_by_tenant(
            app_code=app_code,
            file_id=file_id,
            tenant_id=tenant_id,
            remote_addr=remote_addr
        )
        t = word_reader.read_text_and_tables_from_word(file_path)

        data[f'text_{idx}'] = t

        idx += 1
        file_info = file_info[tail:]

    return data


@bp.route('/load_text_file', methods=['POST', 'GET', ])
def load_text_file():
    parser = reqparse.RequestParser()
    parser.add_argument('filename',
                        location='json', type=str, required=True, )
    args = parser.parse_args()

    filename = args.get('filename')
    with open(f'./file/{filename}', encoding='utf-8', mode='r') as f:
        text = f.read()

    return text


@bp.route('/check_request_json', methods=['POST', 'GET', ])
def check_request_json():
    req_text = request.get_data().decode('utf-8')
    try:
        json.loads(req_text)
        is_json = True
    except Exception:
        is_json = False

    logger.info(f'check_request_json[{is_json}][{req_text}]')

    return {
        'text': req_text,
        'is_json': is_json,
    }


@bp.route('/get_config_text', methods=['POST', 'GET', ])
def get_config_text():
    with open('./file/config_text.txt', encoding='utf-8') as f:
        return f.read()


@bp.route('/issue_coupon', methods=['POST', 'GET', ])
def issue_coupon():
    """
    发券
    :return:
    """
    parser = reqparse.RequestParser()
    parser.add_argument('xxxx',
                        location='json', type=str, required=False, default='xxxx')
    parser.add_argument('workflowVersion',
                        location='json', type=str, required=False, default='1')
    parser.add_argument('mobileNo',
                        location='json', type=str, required=True, )
    parser.add_argument('bankName',
                        location='json', type=str, required=True, )
    args: dict = parser.parse_args()

    workflow_version = args.get('workflowVersion')
    mobile_no = args.get('mobileNo')
    bank_name = args.get('bankName')

    # 测试环境
    base_url = f'https://malla.leagpoint.com/rssz/v1'

    if workflow_version == '1':
        # 测试环境-ZJ测试
        api_key_1 = 'app-5flFXvlgmMer7gOPwg50RswU'
        api_key_2 = 'app-5flFXvlgmMer7gOPwg50RswU'
    elif workflow_version == '2':
        # 测试环境
        # fixme
        api_key_1 = 'app-JabTdjv1Amm8vIKZktURxU8I'
        api_key_2 = 'app-ql3ctWaYqP07EZXfCjMVmBY0'
    elif workflow_version == '3':
        # 生产环境
        # todo
        api_key_1 = 'app-5flFXvlgmMer7gOPwg50RswU'
        api_key_2 = 'app-5flFXvlgmMer7gOPwg50RswU'
        base_url = f'https://dipp.rs-ibg.com/rssz/v1'
    else:
        return {
            'error': f'请求参数无效[workflowVersion][{workflow_version}]',
        }

    logger.info(f'发券-手机号码校验参数[{mobile_no}][{bank_name}]')
    # 手机号码校验
    h = {
        'Authorization': f'Bearer {api_key_1}'
    }
    data = {
        "inputs": {
            'bankName': bank_name,
        },
        'query': mobile_no,
        "response_mode": "blocking",
        "user": str(uuid.uuid4())
    }
    url = f'{base_url}/chat-messages'
    resp = requests.post(url=url, headers=h, json=data, verify=False)
    resp_json: dict = resp.json()
    answer: str = resp_json.get('answer')

    logger.info(f'发券-手机号码校验结果[{mobile_no}][{bank_name}][{answer}]')
    if '发送短信' not in answer:
        # 手机号码校验不通过
        return {
            'text': answer,
        }

    logger.info(f'发券-手机号码校验通过了，可以调用发券了[{mobile_no}][{bank_name}]')

    def do_issue_coupon():
        logger.info(f'发券-调用发券参数[{mobile_no}][{bank_name}]')
        _h = {
            'Authorization': f'Bearer {api_key_2}'
        }
        _data = {
            "inputs": {
                'bankName': bank_name,
            },
            'query': mobile_no,
            "response_mode": "blocking",
            "user": str(uuid.uuid4())
        }
        try:
            _resp = requests.post(url=url, headers=_h, json=_data, verify=False)
            try:
                result = _resp.json()
            except BaseException:
                pass
            result = result if result else _resp.text
            logger.info(f'发券-调用发券结果[{mobile_no}][{bank_name}][{result}]')
        except Exception:
            logger.exception(f'发券-调用发券异常[{mobile_no}][{bank_name}]')

    t = threading.Thread(target=do_issue_coupon, args=())
    t.start()

    return {
            'text': answer,
    }


@bp.route('/get_ocr_result', methods=['POST', 'GET', ])
def get_ocr_result():
    parser = reqparse.RequestParser()
    parser.add_argument('app_code',
                        location='json', type=str, required=True, )
    parser.add_argument('file_info',
                        location='json', type=str, required=True, )
    parser.add_argument('from_addr',
                        location='json', type=str, required=False)
    args = parser.parse_args()

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')

    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'get_ocr_result[{app_code}][{file_info}][{remote_addr}]')

    head = file_info.find("tenant_id='") + len("tenant_id='")
    tail = head + 36
    tenant_id = file_info[head:tail]

    head = file_info.find("related_id='") + len("related_id='")
    tail = head + 36
    file_id = file_info[head:tail]

    file_path = lmjj_agent.file_download_by_tenant(
        app_code=app_code,
        file_id=file_id,
        tenant_id=tenant_id,
        remote_addr=remote_addr
    )
    response_text, texts = paddle_ocr.ocr(file_path)
    logger.info(f'get_ocr_result[{file_path}][{response_text}][{texts}]')
    return {
        "response_text": response_text,
        "texts": texts
    }


@bp.route('/get_ocr_result_v2', methods=['POST', 'GET', ])
def get_ocr_result_v2():
    """
    v2ocr识别，针对长图
    :return:
    """
    parser = reqparse.RequestParser()
    parser.add_argument('app_code',
                        location='json', type=str, required=True, )
    parser.add_argument('file_info',
                        location='json', type=str, required=True, )
    parser.add_argument('from_addr',
                        location='json', type=str, required=False)
    args = parser.parse_args()

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')

    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'get_ocr_result[{app_code}][{file_info}][{remote_addr}]')

    head = file_info.find("tenant_id='") + len("tenant_id='")
    tail = head + 36
    tenant_id = file_info[head:tail]

    head = file_info.find("related_id='") + len("related_id='")
    tail = head + 36
    file_id = file_info[head:tail]

    file_path = lmjj_agent.file_download_by_tenant(
        app_code=app_code,
        file_id=file_id,
        tenant_id=tenant_id,
        remote_addr=remote_addr
    )
    slice_height = 3096
    import cv2
    img = cv2.imread(file_path)
    height, width, _ = img.shape

    # 初始化结果列表
    all_texts = []

    # 计算需要切割的行数
    num_slices = height // slice_height + (height % slice_height > 0)

    for i in range(num_slices):
        # 计算当前切片的起始坐标
        start_y = i * slice_height
        end_y = min((i + 1) * slice_height, height)

        # 切割图片
        sliced_img = img[start_y:end_y, :]
        # 保存或直接使用切片进行OCR，这里以保存为例（实际应用中可能直接处理图像数据）
        try:
            # 使用上下文管理器确保文件在使用完毕后自动关闭和删除
            with tempfile.NamedTemporaryFile(delete=True, suffix=".jpg", mode='w+b') as temp_file:
                cv2.imwrite(temp_file.name, sliced_img)
                temp_file.flush()  # 确保数据写入磁盘

                # 重新定位文件指针到开始位置，以便读取
                temp_file.seek(0)

                response_text, texts = paddle_ocr.ocr(temp_file.name)
                all_texts.extend(texts[0])
                logger.info(f'get_ocr_result[{file_path}][{response_text}][{texts}]')
        except Exception as e:
            logger.error(f'Error processing slice {i}: {e}')
            # 如果发生异常，确保资源被释放，这里由于使用了上下文管理器，文件会被自动关闭和删除

    return {
        "response_text": response_text,
        "texts": [all_texts]
    }

@bp.route('/common_pdf_parse', methods=['POST', 'GET', ])
def common_pdf_parse():
    parser = reqparse.RequestParser()
    parser.add_argument('app_code',
                        location='json', type=str, required=True, )
    parser.add_argument('file_info',
                        location='json', type=str, required=True, )
    parser.add_argument('from_addr',
                        location='json', type=str, required=False)
    args = parser.parse_args()

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')

    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'parse_pdf_table[{app_code}][{file_info}][{remote_addr}]')

    head = file_info.find("tenant_id='") + len("tenant_id='")
    tail = head + 36
    tenant_id = file_info[head:tail]

    head = file_info.find("related_id='") + len("related_id='")
    tail = head + 36
    file_id = file_info[head:tail]

    file_path = lmjj_agent.file_download_by_tenant(
        app_code=app_code,
        file_id=file_id,
        tenant_id=tenant_id,
        remote_addr=remote_addr
    )

    return {'text': pdf_common_parse.extract_information_from_pdf(file_path)}


@bp.route('/convert_to_word', methods=['POST'])
def convert_json_to_word():
    parser = reqparse.RequestParser()
    parser.add_argument('app_code',
                        location='args', type=str, required=True, )
    parser.add_argument('json_data',
                        location='json', type=dict, required=True, )
    parser.add_argument('from_addr',
                        location='json', type=str, required=False)
    args = parser.parse_args()
    app_code = args.get('app_code')
    word_data: dict = args.get('json_data')["data"]
    from_addr = args.get('from_addr')
    try:
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = '黑体'
        rFonts = style._element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), '黑体')
        document.add_heading('广告用语审查结果', 0)
        for index, item in enumerate(word_data, start=1):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"{index}. 问题词汇：")
            run.bold = True
            run.font.size = Pt(12)
            run.add_text(item['problemWord'])
            document.add_paragraph('问题类型：' + item['problemType'])
            document.add_paragraph('相关法律：' + item['laws'])
            document.add_paragraph('审查意见：' + item['reviewComment'])
        tmp_filename = f'消保审查-{str(uuid.uuid4())}.docx'
        tmp_filepath = f'{tempfile.gettempdir()}/{tmp_filename}'
        try:
            document.save(tmp_filepath)
        except Exception:
            logger.error(f'临时文件生成失败[{tmp_filepath}]', exc_info=True)
            return {
                'error': f'临时文件生成失败',
            }

        result = lmjj_agent.upload_file(app_code=app_code, file=tmp_filepath, remote_addr=from_addr)
        result['error'] = ''
        return result

    except Exception:
        logger.error(f'word生成失败', exc_info=True)
        return {
            'error': 'word生成失败',
        }


@bp.route('/read_excel', methods=['POST', 'GET', ])
def read_excel():
    req_text = request.data.decode('utf-8')
    logger.info(
        f'read_excel'
        f'[{req_text}]'
    )
    try:
        args: dict = json.loads(req_text)
    except Exception:
        logger.error(
            f'read_excel'
            f'[请求参数不是正确的JSON]',
            exc_info=True
        )
        return {
            'error': '请求参数不是正确的JSON',
            'text': '',
        }

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')

    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'read_word_v4[{app_code}][{file_info}][{remote_addr}]')

    # 目前最多只能上传[6]个文件
    # 预分配参数防止工作流crash
    data = {
        'text_1': '',
        'text_2': '',
        'text_3': '',
        'text_4': '',
        'text_5': '',
        'text_6': '',
    }
    idx = 1
    while True:
        head = file_info.find("tenant_id='")
        if head < 0:
            break
        head += + len("tenant_id='")
        tail = head + 36
        tenant_id = file_info[head:tail]

        head = file_info.find("related_id='")
        head += + len("related_id='")
        tail = head + 36
        file_id = file_info[head:tail]

        file_path = lmjj_agent.file_download_by_tenant(
            app_code=app_code,
            file_id=file_id,
            tenant_id=tenant_id,
            remote_addr=remote_addr
        )
        t = excel_to_json.read_excel(file_path)

        data[f'text_{idx}'] = t

        idx += 1
        file_info = file_info[tail:]

    return data


@bp.route('/semantic_label', methods=['POST'])
async def semantic_label():
    """
    语义理解标签
    :return:
    """
    req_text = request.data.decode('utf-8')
    logger.info(
        f'semantic_label'
        f'[{req_text}]'
    )
    try:
        args: dict = json.loads(req_text)
    except Exception:
        logger.error(
            f'semantic_label'
            f'[请求参数不是正确的JSON]',
            exc_info=True
        )
        return {
            'error': '请求参数不是正确的JSON',
            'text': '',
        }

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')
    output_len: int = args.get('output_len')
    sleep_time: int = args.get('sleep_time')
    app_key: str = args.get('app_key')
    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'semantic_label[{app_code}][{file_info}][{remote_addr}][{output_len}][{sleep_time}][{app_key}]')
    base_url = from_addr + "v1"
    headers = {'Authorization': f'Bearer {app_key}'}
    url = f'{base_url}/workflows/run'
    logger.info(f'外呼语义标签,  api_key:{app_key}, base_url:{base_url}')
    # 目前最多只能上传[6]个文件
    # 预分配参数防止工作流crash
    # download_urls = {
    #     'url_1': '',
    #     'url_2': '',
    #     'url_3': '',
    #     'url_4': '',
    #     'url_5': '',
    #     'url_6': '',
    # }

    df_list = []
    while True:
        head = file_info.find("tenant_id='")
        if head < 0:
            break
        head += + len("tenant_id='")
        tail = head + 36
        tenant_id = file_info[head:tail]

        head = file_info.find("related_id='")
        head += + len("related_id='")
        tail = head + 36
        file_id = file_info[head:tail]

        file_path = lmjj_agent.file_download_by_tenant(
            app_code=app_code,
            file_id=file_id,
            tenant_id=tenant_id,
            remote_addr=remote_addr
        )
        df = pd.read_excel(file_path)
        df = df[df["解析命令"].notna()]
        df = df[df["解析命令"] != "nomatch:out-of-voca"]
        # df["解析命令"]这一列 删除 空值 "nomatch:out-of-voca" 的行
        # df["解析命令"] = df["解析命令"].replace(np.nan, "无声音")
        # df["解析命令"] = df["解析命令"].replace("nomatch:out-of-voca", "无识别")
        df_list.append(df)
        file_info = file_info[tail:]
    # 只解析第一个文件
    df = df_list[0]
    # 定义异步请求函数
    async def send_post_request(session, url, headers, data):
        async with session.post(url=url, headers=headers, json=data, ssl=True) as response:
            if response.status != 200:
                logger.error(f"Failed to post to {url} with status code {response.status},{response.text}")
                return ""
            return await response.json()
    # 创建异步任务列表
    tasks = []
    # 切分数据
    df1_slice_list = [df.iloc[i:i+output_len] for i in range(0,len(df),output_len)]
    # 使用同一个session来复用TCP连接
    # df1_slice_list = df1_slice_list[:1]
    async with ClientSession() as session:
        # 将 df 解析命令  列取出，使用zip 与1-100序号配对后输出字符串
        for index, df1_slice in enumerate(df1_slice_list):
            promopt_str_list = [f"{i}.{j}" for i, j in zip(range(1, output_len + 1), df1_slice["解析命令"])]
            promopt_str = "\n".join(promopt_str_list)
            data = {
                "inputs": {
                    "text": promopt_str
                },
                "response_mode": "blocking",
                "user": str(uuid.uuid4())
            }
            logger.info(f'异步请求第{index+1}个')
            task = asyncio.ensure_future(send_post_request(session, url, headers, data))
            tasks.append(task)
            time.sleep(sleep_time)
            # if (index+1) % 10 == 0:
            #     time.sleep(120)

        # 等待所有请求完成
        responses = await asyncio.gather(*tasks)

    # 汇总结果
    result = {'responses': responses}
    logger.info(f'异步请求结果:{result}')
    # 使用pd解析ret的text
    ret_text_list = [i["data"]["outputs"]["text"] for i in result["responses"]]
    ret_df_list = [pd.read_csv(io.StringIO(i), sep=",", header=None) for i in ret_text_list]
    # 修改列名，添加序号，并且给序号列添加前缀
    ret_df_list = [i.rename(columns={0: "序号", 1: "大模型-解析命令"}) for i in ret_df_list]
    # 给序号列的元素添加前缀,前缀为 在ret_df_list中的位置
    for j, i in enumerate(ret_df_list):
        i["序号"] = i["序号"].apply(lambda x: f"{j}-{x}")
        i["大模型-解析命令"] = i["大模型-解析命令"].apply(lambda x: x.strip())
    # 把ret_df_list合并成一个df
    ret_df2 = pd.concat(ret_df_list)
    # 把df1_slice_list合并成一个
    ret_df1 = pd.concat(df1_slice_list)
    # 把ret_df2和ret_df1的index列重置
    ret_df2.reset_index(drop=True, inplace=True)
    ret_df1 = ret_df1.reset_index(drop=True)
    # 把ret_df2和ret_df1合并成一个
    ret_df = pd.concat([ret_df1, ret_df2], axis=1)
    tmp_filename = f'识别对话详情-大模型-{str(uuid.uuid4())}.xlsx'
    tmp_filepath = f'{tempfile.gettempdir()}/{tmp_filename}'
    try:
        #保存ret_df，防止中文乱码
        ret_df.to_excel(tmp_filepath, index=False)
    except Exception:
        logger.error(f'临时文件生成失败[{tmp_filepath}]', exc_info=True)
        return {
            'error': f'临时文件生成失败',
        }
    download_result = lmjj_agent.upload_file(app_code=app_code, file=tmp_filepath, remote_addr=from_addr)
    download_result['error'] = ''
    return download_result


@bp.route('/semantic_label_v2', methods=['POST'])
async def semantic_label_v2():
    """
    todo：语义理解标签v2 带上客服话术
    :return:
    """
    req_text = request.data.decode('utf-8')
    logger.info(
        f'semantic_label'
        f'[{req_text}]'
    )
    try:
        args: dict = json.loads(req_text)
    except Exception:
        logger.error(
            f'semantic_label'
            f'[请求参数不是正确的JSON]',
            exc_info=True
        )
        return {
            'error': '请求参数不是正确的JSON',
            'text': '',
        }

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')
    output_len: int = args.get('output_len')
    sleep_time: int = args.get('sleep_time')
    app_key: str = args.get('app_key')
    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'semantic_label[{app_code}][{file_info}][{remote_addr}][{output_len}][{sleep_time}][{app_key}]')
    base_url = from_addr + "v1"
    headers = {'Authorization': f'Bearer {app_key}'}
    url = f'{base_url}/workflows/run'
    logger.info(f'外呼语义标签,  api_key:{app_key}, base_url:{base_url}')
    df_list = []
    while True:
        head = file_info.find("tenant_id='")
        if head < 0:
            break
        head += + len("tenant_id='")
        tail = head + 36
        tenant_id = file_info[head:tail]

        head = file_info.find("related_id='")
        head += + len("related_id='")
        tail = head + 36
        file_id = file_info[head:tail]

        file_path = lmjj_agent.file_download_by_tenant(
            app_code=app_code,
            file_id=file_id,
            tenant_id=tenant_id,
            remote_addr=remote_addr
        )
        df = pd.read_excel(file_path)
        # df = df[df["解析命令"].notna()]
        # df = df[df["解析命令"] != "nomatch:out-of-voca"]
        # df["解析命令"]这一列 删除 空值 "nomatch:out-of-voca" 的行
        df["解析命令"] = df["解析命令"].replace(np.nan, "无声音")
        df["解析命令"] = df["解析命令"].replace("nomatch:out-of-voca", "无识别")

        df_list.append(df)
        file_info = file_info[tail:]
    # 只解析第一个文件
    df = df_list[0]
    # 定义异步请求函数
    async def send_post_request(session, url, headers, data):
        async with session.post(url=url, headers=headers, json=data, ssl=True) as response:
            if response.status != 200:
                logger.error(f"Failed to post to {url} with status code {response.status},{response.text}")
                return ""
            return await response.json()
    # 创建异步任务列表
    tasks = []
    # 切分数据
    df1_slice_list = [df.iloc[i:i+output_len] for i in range(0,len(df),output_len)]
    # 使用同一个session来复用TCP连接
    # df1_slice_list = df1_slice_list[:1]
    async with ClientSession() as session:
        # 将 df 解析命令  列取出，使用zip 与1-100序号配对后输出字符串
        for index, df1_slice in enumerate(df1_slice_list):
            promopt_str_list = [f"{i}.{j}" for i, j in zip(range(1, output_len + 1), df1_slice["解析命令"])]
            promopt_str = "\n".join(promopt_str_list)
            data = {
                "inputs": {
                    "text": promopt_str
                },
                "response_mode": "blocking",
                "user": str(uuid.uuid4())
            }
            logger.info(f'异步请求第{index+1}个')
            task = asyncio.ensure_future(send_post_request(session, url, headers, data))
            tasks.append(task)
            time.sleep(sleep_time)
            # if (index+1) % 10 == 0:
            #     time.sleep(120)

        # 等待所有请求完成
        responses = await asyncio.gather(*tasks)

    # 汇总结果
    result = {'responses': responses}
    logger.info(f'异步请求结果:{result}')
    # 使用pd解析ret的text
    ret_text_list = [i["data"]["outputs"]["text"] for i in result["responses"]]
    ret_df_list = [pd.read_csv(io.StringIO(i), sep=",", header=None) for i in ret_text_list]
    # 修改列名，添加序号，并且给序号列添加前缀
    ret_df_list = [i.rename(columns={0: "序号", 1: "大模型-解析命令"}) for i in ret_df_list]
    # 给序号列的元素添加前缀,前缀为 在ret_df_list中的位置
    for j, i in enumerate(ret_df_list):
        i["序号"] = i["序号"].apply(lambda x: f"{j}-{x}")
    # 把ret_df_list合并成一个df
    ret_df2 = pd.concat(ret_df_list)
    # 把df1_slice_list合并成一个
    ret_df1 = pd.concat(df1_slice_list)
    # 把ret_df2和ret_df1的index列重置
    ret_df2.reset_index(drop=True, inplace=True)
    ret_df1 = ret_df1.reset_index(drop=True)
    # 把ret_df2和ret_df1合并成一个
    ret_df = pd.concat([ret_df1, ret_df2], axis=1)
    tmp_filename = f'识别对话详情-大模型-{str(uuid.uuid4())}.xlsx'
    tmp_filepath = f'{tempfile.gettempdir()}/{tmp_filename}'
    try:
        #保存ret_df，防止中文乱码
        ret_df.to_excel(tmp_filepath, index=False)
    except Exception:
        logger.error(f'临时文件生成失败[{tmp_filepath}]', exc_info=True)
        return {
            'error': f'临时文件生成失败',
        }
    download_result = lmjj_agent.upload_file(app_code=app_code, file=tmp_filepath, remote_addr=from_addr)
    download_result['error'] = ''
    return download_result


@bp.route('/chat_summary', methods=['POST'])
async def chat_summary():
    """
    语义理解标签
    :return:
    """
    req_text = request.data.decode('utf-8')
    logger.info(
        f'chat_summary'
        f'[{req_text}]'
    )
    try:
        args: dict = json.loads(req_text)
    except Exception:
        logger.error(
            f'chat_summary'
            f'[请求参数不是正确的JSON]',
            exc_info=True
        )
        return {
            'error': '请求参数不是正确的JSON',
            'text': '',
        }

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')
    output_len: int = args.get('output_len')
    sleep_time: int = args.get('sleep_time')
    app_key: str = args.get('app_key')

    remote_addr: str = from_addr if from_addr else request.remote_addr
    base_url = from_addr + "v1"
    headers = {'Authorization': f'Bearer {app_key}'}
    url = f'{base_url}/workflows/run'
    logger.info(f'对话总结请求, app_code:{app_code}, file_info:{file_info}, from_addr:{from_addr}, output_len:{output_len},'
                f'sleep_time:{sleep_time},  api_key:{app_key}, base_url:{base_url}')
    # 目前最多只能上传[6]个文件
    # 预分配参数防止工作流crash
    # download_urls = {
    #     'url_1': '',
    #     'url_2': '',
    #     'url_3': '',
    #     'url_4': '',
    #     'url_5': '',
    #     'url_6': '',
    # }

    df_list = []
    while True:
        head = file_info.find("tenant_id='")
        if head < 0:
            break
        head += + len("tenant_id='")
        tail = head + 36
        tenant_id = file_info[head:tail]

        head = file_info.find("related_id='")
        head += + len("related_id='")
        tail = head + 36
        file_id = file_info[head:tail]

        file_path = lmjj_agent.file_download_by_tenant(
            app_code=app_code,
            file_id=file_id,
            tenant_id=tenant_id,
            remote_addr=remote_addr
        )
        df = pd.read_excel(file_path)
        df["解析命令"] = df["解析命令"].replace(np.nan, "无声音")
        df["解析命令"] = df["解析命令"].replace("nomatch:out-of-voca", "无识别")
        df_list.append(df)
        file_info = file_info[tail:]
    # 只解析第一个文件
    df = df_list[0]
    # 定义异步请求函数
    async def send_post_request(session, url, headers, data):
        try:
            async with session.post(url=url, headers=headers, json=data, ssl=True) as response:
                if response.status != 200:
                    logger.error(f"Failed to post to {url} with status code {response.status},{response.text}")
                    return {
                        "data":
                            {"outputs":
                                 {"result": "本次总结请求出错"}}}
                return await response.json()
        except asyncio.exceptions.TimeoutError:
            logger.error(f"Failed to post to {url} with status code {response.status},{response.text}")
            return {
                "data":
                    {"outputs":
                         {"result": "本次总结请求出错"}}}
    # 创建异步任务列表
    tasks = []
    # 切分一次需要的数据
    case_num_list = df["案件编号"].unique()
    # 用前几个测试
    # case_num_list = case_num_list[:50]
    promopt_str_dict = {}
    for case_id in case_num_list:
        extracted_data = df[df['案件编号'] == case_id][['AI话术', '解析命令']]
        # 将提取的数据转换为字符串格式，每条记录一行
        output_strings = extracted_data.apply(lambda row: f"客服: {row['AI话术']}，客户: {row['解析命令']}",
                                              axis=1).tolist()
        promopt_str = "\n".join(output_strings)
        promopt_str_dict[case_id] = promopt_str
    # 使用同一个session来复用TCP连接
    # df1_slice_list = df1_slice_list[:1]
    promopt_str_dict_values = [promopt_str_dict[case_id] for case_id in case_num_list]

    async with ClientSession() as session:
        # 将 df 解析命令  列取出，使用zip 与1-100序号配对后输出字符串
        for index, promopt_str in enumerate(promopt_str_dict_values):
            data = {
                "inputs": {
                    "connectCot": promopt_str
                },
                "response_mode": "blocking",
                "user": str(uuid.uuid4())
            }
            logger.info(f'异步请求第{index+1}个')
            task = asyncio.ensure_future(send_post_request(session, url, headers, data))
            tasks.append(task)
            time.sleep(sleep_time)
            # if (index+1) % 10 == 0:
            #     time.sleep(120)

        # 等待所有请求完成
        responses = await asyncio.gather(*tasks)

    # 汇总结果
    result = {'responses': responses}
    logger.info(f'异步请求结果:{result}')
    ret_text_list = []
    for i in result["responses"]:
        try:
            ret_text_list.append(i["data"]["outputs"]["result"])
        except Exception:
            ret_text_list.append("本次总结请求出错")
    # 给df增加一列，名为总结的空列
    df["总结"] = [None] * len(df)
    # 把ret_text_list 按照case_num_list的顺序 依据case_id 添加到对应的df["总结"]列
    for i, case_id in enumerate(case_num_list):
        df.loc[df["案件编号"] == case_id, "总结"] = ret_text_list[i]
    tmp_filename = f'识别对话总结-{str(uuid.uuid4())}.xlsx'
    tmp_filepath = f'{tempfile.gettempdir()}/{tmp_filename}'

    try:
        #保存ret_df，防止中文乱码
        df.to_excel(tmp_filepath, index=False)
        df.to_excel(f"./file/{tmp_filename}", index=False, encoding='utf-8')
    except Exception:
        logger.error(f'临时文件生成失败[{tmp_filepath}]', exc_info=True)
        return {
            'error': f'临时文件生成失败',
        }
    download_result = lmjj_agent.upload_file(app_code=app_code, file=tmp_filepath, remote_addr=from_addr)
    download_result['error'] = ''
    return download_result

@bp.route('/activate_sleeping_users', methods=['POST'])
async def activate_sleeping_users():
    '''
    用户促活
    '''
    req_text = request.data.decode('utf-8')
    path = get_file_path.get_path(req_text)
    parser = reqparse.RequestParser()
    parser.add_argument('app_code',
                        location='json', type=str, required=True, )
    parser.add_argument('app_key',
                        location='json', type=str, required=True, )
    parser.add_argument('file_info',
                        location='json', type=str, required=True, )
    parser.add_argument('from_addr',
                        location='json', type=str, required=False)
    parser.add_argument('sleep_time',
                        location='json', type=int, required=False)
    args = parser.parse_args()

    app_code: str = args.get('app_code')
    file_info: str = args.get('file_info')
    from_addr: str = args.get('from_addr')
    app_key: str = args.get('app_key')
    sleep_time: int = args.get('sleep_time')
    base_url = from_addr + "v1"
    headers = {'Authorization': f'Bearer {app_key}'}
    url = f'{base_url}/chat-messages'
    remote_addr: str = from_addr if from_addr else request.remote_addr
    logger.info(f'用户促活请求:[{app_code}][{file_info}][{remote_addr}]')
    # 只上传一个文件
    excel_path = path["path_1"]
    df = pd.read_excel(excel_path)

    # 创建异步任务列表
    tasks = []
    async with ClientSession() as session:
        # 将 df 解析命令  列取出，使用zip 与1-100序号配对后输出字符串
        for index, df_row_i in enumerate(df.values):
            data = {
                "inputs": {
                    "info": f"手机号：{df_row_i[0]},姓名：{df_row_i[1]}",
                    "bill": f"{df_row_i[2]}"
                },
                "query": "1",
                "response_mode": "blocking",
                "conversation_id": "",
                "user": str(uuid.uuid4())
            }
            logger.info(f'异步请求第{index + 1}个,{data}')
            task = asyncio.ensure_future(async_request.send_post_request(session, url, headers, data))
            tasks.append(task)
            time.sleep(sleep_time)
            # if (index+1) % 10 == 0:
            #     time.sleep(120)

            # 等待所有请求完成
        responses = await asyncio.gather(*tasks)
    # 汇总结果
    result = {'responses': responses}
    logger.info(f'异步请求结果:{result}')
    ret_text_list = []
    for i in result["responses"]:
        try:
            ret_text_list.append(i["answer"])
        except Exception:
            ret_text_list.append("本次总结请求出错")
    # 给df增加一列，名为结果空列
    df["结果"] = [None] * len(df)
    # 把ret_text_list 按照case_num_list的顺序 依据case_id 添加到对应的df["总结"]列
    for i, i_result in enumerate(ret_text_list):
        df.loc[i, "结果"] = i_result
    tmp_filename = f'用户促活结果-{str(uuid.uuid4())}.xlsx'
    tmp_filepath = f'{tempfile.gettempdir()}/{tmp_filename}'
    try:
        # 保存ret_df，防止中文乱码
        df.to_excel(tmp_filepath, index=False)
    except Exception:
        logger.error(f'临时文件生成失败[{tmp_filepath}]', exc_info=True)
        return {
            'error': f'临时文件生成失败',
        }
    download_result = lmjj_agent.upload_file(app_code=app_code, file=tmp_filepath, remote_addr=from_addr)
    download_result['error'] = ''
    return download_result



# 调度器管理API
@bp.route('/scheduler/status', methods=['GET'])
def scheduler_status():
    """获取调度器状态"""
    jobs = []
    for job in scheduler.get_jobs():
        jobs.append({
            'id': job.id,
            'name': job.func.__name__,
            'args': str(job.args),
            'next_run_time': str(job.next_run_time),
            'trigger': str(job.trigger)
        })
    
    return {
        'status': 'running' if scheduler.running else 'stopped',
        'job_count': len(scheduler.get_jobs()),
        'jobs': jobs
    }

@bp.route('/scheduler/start', methods=['POST'])
def scheduler_start():
    """启动调度器"""
    global scheduler
    
    if not scheduler or not hasattr(scheduler, 'running') or not scheduler.running:
        # 检查已有调度器
        from apscheduler.schedulers.background import BackgroundScheduler
        
        # 获取所有实例中的第一个（如果有）
        import gc
        schedulers = [obj for obj in gc.get_objects() if isinstance(obj, BackgroundScheduler)]
        
        if schedulers:
            # 使用已存在的调度器
            return {'status': 'warning', 'message': '调度器已在其他地方初始化，请使用已有调度器'}
        else:
            # 导入需要的函数
            from scheduler import init_scheduler
            init_scheduler()
            return {'status': 'success', 'message': '调度器已初始化并启动'}
    else:
        return {'status': 'warning', 'message': '调度器已在运行中'}

@bp.route('/scheduler/stop', methods=['POST'])
def scheduler_stop():
    """停止调度器"""
    # 导入shutdown_scheduler函数
    from scheduler import shutdown_scheduler
    
    if scheduler and scheduler.running:
        shutdown_scheduler()
        return {'status': 'success', 'message': '调度器已停止'}
    else:
        return {'status': 'warning', 'message': '调度器未运行'}

@bp.route('/scheduler/run_task', methods=['POST'])
def run_task():
    """立即运行指定任务"""
    parser = reqparse.RequestParser()
    parser.add_argument('task_name', location='json', type=str, required=True)
    parser.add_argument('args', location='json', type=list, required=False)
    args = parser.parse_args()
    
    # 导入API token
    from scheduler import WECHAT_API_TOKEN, API_TOKEN
    
    task_name = args.get('task_name')
    task_args = args.get('args', [])
    
    if task_name == 'hotspots' and len(task_args) > 0:
        # 根据任务类型选择不同的API令牌
        original_tasks = ["每日金句", "金价", "金融", "房地产", "经济指标"]
        if task_args[0] in original_tasks:
            # 使用原始任务的API_TOKEN
            api_token = task_args[1] if len(task_args) > 1 else API_TOKEN
        else:
            # 使用其他任务的WECHAT_API_TOKEN
            api_token = task_args[1] if len(task_args) > 1 else WECHAT_API_TOKEN
        TaskConfig.hotspots(task_args[0], api_token)
        return {'status': 'success', 'message': f'已执行任务 {task_name}，参数：{task_args}'}
    elif task_name == 'hourly_task' and len(task_args) > 0:
        # 执行hourly_task任务
        api_token = task_args[1] if len(task_args) > 1 else WECHAT_API_TOKEN
        TaskConfig.hourly_task(task_args[0], api_token)
        return {'status': 'success', 'message': f'已执行小时任务，参数：{task_args}'}
    elif task_name == 'task2':
        TaskConfig.task2(API_TOKEN)
        return {'status': 'success', 'message': f'已执行任务 {task_name}'}
    else:
        return {'status': 'error', 'message': f'未找到任务 {task_name} 或参数不正确'}
